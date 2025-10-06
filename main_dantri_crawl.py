import time
import os
from selenium.webdriver.common.by import By
import undetected_chromedriver as uc
from docx import Document

# HÀM KHỞI TẠO DRIVER (CHẾ ĐỘ STEALTH)
def create_driver():
    options = uc.ChromeOptions()
    options.add_argument("--headless=new")  # bỏ nếu muốn thấy trình duyệt
    options.add_argument("--no-sandbox")
    options.add_argument("--disable-dev-shm-usage")
    options.add_argument("--disable-blink-features=AutomationControlled")
    options.add_argument("--disable-popup-blocking")
    options.add_argument("--start-maximized")
    options.add_argument("--disable-infobars")
    options.add_argument(
        "--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/130.0.0.0 Safari/537.36"
    )

    driver = uc.Chrome(options=options)
    driver.delete_all_cookies()
    return driver


# LẤY DANH SÁCH LINK CÁC BÀI TRONG TRANG (có fallback selector)
def get_article_links(driver, url):
    driver.get(url)
    time.sleep(2)

    # Thử selector chính xác hơn trước, nếu không có thì fallback
    links = []
    try:
        elems = driver.find_elements(By.CSS_SELECTOR, "h3.article-title a[href]")
        for a in elems:
            href = a.get_attribute("href")
            if href and href.startswith("https://dantri.com.vn/thoi-su/") and href not in links:
                links.append(href)
    except Exception:
        pass

    # Fallback: chung chung nếu không tìm thấy với selector trên
    if not links:
        try:
            elems = driver.find_elements(By.CSS_SELECTOR, "article a[href]")
            for a in elems:
                href = a.get_attribute("href")
                if href and href.startswith("https://dantri.com.vn/thoi-su/") and href not in links:
                    links.append(href)
        except Exception:
            pass

    print(f"📰 Trang {url} → Tìm thấy {len(links)} bài")
    return links


# LẤY NỘI DUNG TỪ MỘT LINK
def get_article_content(driver, url):
    try:
        driver.get(url)
        time.sleep(2)
        title = driver.find_element(By.CSS_SELECTOR, "h1").text.strip()

        # Ngày đăng (nếu có)
        try:
            date = driver.find_element(By.CSS_SELECTOR, "time").text.strip()
        except:
            date = "Không rõ ngày đăng"

        # Nội dung bài viết (selector chung)
        paragraphs = driver.find_elements(By.CSS_SELECTOR, "div.dt-news__content p, div.singular-content p, article p")
        content = "\n".join([p.text for p in paragraphs if p.text.strip()])

        if not content:
            raise ValueError("Không có nội dung")

        return {"title": title, "date": date, "content": content, "url": url}

    except Exception as e:
        print(f"⚠️ Lỗi khi đọc bài {url}: {e}")
        return None


# LƯU VÀO WORD (THÊM DẦN VÀO FILE ĐÃ TỒN TẠI)
def save_to_word(articles, file_path=r"F:\KPDL_HLV\Final.docx"):
    if not os.path.exists(file_path):
        doc = Document()
        doc.save(file_path)

    try:
        doc = Document(file_path)
    except:
        doc = Document()

    for art in articles:
        doc.add_heading(art["title"], level=1)
        doc.add_paragraph(f"🕓 Ngày đăng: {art['date']}")
        doc.add_paragraph(art["content"])
        doc.add_paragraph(f"🔗 Nguồn: {art['url']}")
        doc.add_paragraph("\n" + "-" * 80 + "\n")

    doc.save(file_path)
    print(f"💾 Đã thêm {len(articles)} bài vào {file_path}")


# MAIN (crawl tuần tự trang 1 -> 6, dùng đúng đường dẫn /trang-X.htm)
def main():
    base_url = "https://dantri.com.vn/thoi-su/nong-tren-mang"  # không có .htm ở đây để dễ ghép
    driver = create_driver()

    total_saved = 0

    try:
        # Crawl tuần tự từng trang (1 → 6)
        for page_num in range(1, 7):
            if page_num == 1:
                page_url = f"{base_url}.htm"
            else:
                page_url = f"{base_url}/trang-{page_num}.htm"

            print(f"\n🌐 Đang xử lý trang {page_num}: {page_url}")

            article_links = get_article_links(driver, page_url)
            page_articles = []

            for i, link in enumerate(article_links, start=1):
                art = get_article_content(driver, link)
                if art:
                    page_articles.append(art)
                    print(f"✅ ({i}/{len(article_links)}) {art['title'][:60]}...")
                else:
                    print(f"⚠️ Bỏ qua bài lỗi: {link}")

                # nghỉ nhẹ giữa các bài để tránh bị chặn
                time.sleep(1.2)

            # Lưu các bài của trang hiện tại vào Word
            if page_articles:
                save_to_word(page_articles)
                total_saved += len(page_articles)
                print(f"📘 Trang {page_num} đã lưu {len(page_articles)} bài.\n")

            # nghỉ nhẹ giữa các trang
            time.sleep(2)

    finally:
        try:
            driver.quit()
        except Exception:
            pass

    print(f"\n🎉 Hoàn tất! Tổng cộng đã lưu {total_saved} bài vào Final.docx")


if __name__ == "__main__":
    main()
